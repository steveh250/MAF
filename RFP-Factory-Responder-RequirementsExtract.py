## Author: Steve Harris
# Purpose: MAF Agent with RAG that checks Company Info against an RFP
# NOTES:
#  - Requires: pip install chromadb python-docx

import asyncio
import uuid
import sys
import chromadb
from chromadb.utils import embedding_functions
from agent_framework import ChatAgent, MCPStreamableHTTPTool
from agent_framework.openai import OpenAIChatClient
from docx import Document
from datetime import datetime
import logging
import os, re, json

# Suppress noisy MCP async generator cleanup warnings
logging.getLogger('asyncio').setLevel(logging.CRITICAL)

# --- RAG Manager Class Implementation ---

class RAGManager:
    def __init__(self):
        self.client = chromadb.Client()
        self.ef = embedding_functions.DefaultEmbeddingFunction()
        self.collection = self.client.get_or_create_collection(
            name="rfp_knowledge",
            embedding_function=self.ef
        )

    def reset_database(self):
        """Clears existing data to ensure fresh context."""
        try:
            self.client.delete_collection("rfp_knowledge")
        except:
            pass
        self.collection = self.client.get_or_create_collection(
            name="rfp_knowledge",
            embedding_function=self.ef
        )
        print(f"Database reset successfully.\n")
        return "Database reset successfully."

    def add_document(self, content: str, source: str) -> str:
        """
        Adds text content to the vector database.
        Args:
            content: The text (markdown) to add.
            source: A label for where this text came from (e.g., "Company Profile", "RFP").
        """

        name="add_to_knowledge_base",
        description="Saves text to the vector DB. Requires 'content' and 'source' (e.g., 'Company-Info' or 'RFP').",

        # Split by double newline to create logical chunks

        chunks = [c for c in content.split("\n\n") if c.strip()]

        if not chunks:
            return "No content found to add."

        ids = [str(uuid.uuid4()) for _ in chunks]
        metadatas = [{"source": source} for _ in chunks]

        self.collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        print(f"--- [RAG] Ingested {len(chunks)} chunks from source: {source} ---")
        return f"Successfully added {len(chunks)} chunks from '{source}' to the knowledge base."

    def query_knowledge(self, query: str, n_results: int = 10) -> str:
        """
        Queries the vector database for relevant context.
        """

        name="query_knowledge_base",
        description="Searches the vector DB for answers. Useful for matching RFP requirements to Company capabilities.",

        results = self.collection.query(
            query_texts=[query],
            n_results=n_results
        )

        context = ""
        if results['documents']:
            for i, (doc, meta) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
                source_label = meta.get('source', 'Unknown')
                context += f"-- Result {i+1} (Source: {source_label}) --\n{doc}\n\n"

        print(f"Knowledge query executed: '{query}'\n")
        return context if context else "No relevant information found in the knowledge base."


def save_to_json(result, output_filename: str = "rfp_response.json"):
    """Extracts JSON payload for output file and saves model reasoning trace separately for audit."""

    # Extract text from Agent response
    if hasattr(result, 'text'):
        content = result.text
    elif hasattr(result, 'content'):
        content = result.content
    else:
        content = str(result)

    # Ensure audit filename ends in .audit.json
    base, ext = os.path.splitext(output_filename)
    audit_filename = f"{base}.audit.json"

    # --- Extract <think> reasoning content for audit log ---
    think_matches = re.findall(r"<think>(.*?)</think>", content, flags=re.DOTALL)
    reasoning_log = [t.strip() for t in think_matches if t.strip()]

    # --- Remove <think> blocks from main output content ---
    content_no_think = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()

    # --- Extract JSON fenced block if present ---
    json_str = None
    fence = "```json"
    if fence in content_no_think:
        start = content_no_think.index(fence) + len(fence)
        end = content_no_think.find("```", start)
        if end == -1:
            end = len(content_no_think)
        json_str = content_no_think[start:end].strip()
    else:
        # fallback: find braces
        first_brace = content_no_think.find("{")
        last_brace = content_no_think.rfind("}")
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_str = content_no_think[first_brace:last_brace + 1].strip()

    # --- Parse JSON or fallback ---
    if json_str:
        try:
            data = json.loads(json_str)
        except Exception:
            data = {
                "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "raw_output": content_no_think
            }
    else:
        data = {
            "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "raw_output": content_no_think
        }

    # --- Write main output JSON ---
    with open(output_filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # --- Write reasoning audit log ---
    audit_record = {
        "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "reasoning_trace": reasoning_log,
        "raw_full_response": content  # full record for evidence
    }
    with open(audit_filename, "w", encoding="utf-8") as f:
        json.dump(audit_record, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Extracted JSON saved to: {output_filename}")
    print(f"✓ Full audit + reasoning saved to: {audit_filename}")

    return output_filename, audit_filename


# --- Helper: strip <think> blocks for per-section answers ---

def extract_main_content(result) -> str:
    """
    Returns the model's main content with <think> blocks stripped out.
    This is used for the per-section answer agent.
    """
    if hasattr(result, 'text'):
        content = result.text
    elif hasattr(result, 'content'):
        content = result.content
    else:
        content = str(result)

    content_no_think = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
    return content_no_think


# --- Helper: iterate answerable nodes (bottom-up within each dict) ---

def iter_answerable_nodes(node, path=None):
    """
    Recursively walks the RFP requirements JSON and yields
    answerable nodes for ALL layers of the JSON.

    Rules:
      - If we see a dict: iterate keys in reverse order (bottom-up within section).
      - If we see a list of primitives: treat the list as a single answerable node.
      - If we see a list of dicts/lists: recurse into each item.
      - If we see a single primitive (string/number/bool/None): treat it as
        an answerable node with a single requirement.

    Yields dicts with:
      - path: list of keys/indexes from root to this node
      - heading: the "label" for this node (usually the last key or index)
      - requirements: list of strings (the bullets or text to respond to)
    """
    if path is None:
        path = []

    # Helper to decide if something is a "primitive"
    def is_primitive(x):
        return isinstance(x, (str, int, float, bool)) or x is None

    # --- Case 1: dict ---
    if isinstance(node, dict):
        # Bottom-up within each section
        for key in reversed(list(node.keys())):
            value = node[key]
            new_path = path + [key]

            # If the value is a list, decide whether it's a leaf or container
            if isinstance(value, list):
                if all(is_primitive(v) for v in value):
                    # Leaf: list of primitive requirements
                    yield {
                        "path": new_path,
                        "heading": key,
                        "requirements": [str(v) for v in value]
                    }
                else:
                    # Container: recurse into each element
                    for idx, item in enumerate(value):
                        item_path = new_path + [f"[{idx}]"]
                        yield from iter_answerable_nodes(item, item_path)

            # If the value is another dict, recurse into it
            elif isinstance(value, dict):
                yield from iter_answerable_nodes(value, new_path)

            # If the value is a single primitive, treat it as a leaf requirement
            elif is_primitive(value):
                yield {
                    "path": new_path,
                    "heading": key,
                    "requirements": [str(value)]
                }

    # --- Case 2: list at this level ---
    elif isinstance(node, list):
        if node and all(is_primitive(v) for v in node):
            # Leaf list of primitives
            heading = path[-1] if path else "[root_list]"
            yield {
                "path": path,
                "heading": heading,
                "requirements": [str(v) for v in node]
            }
        else:
            # Container list: recurse into children
            for idx, item in enumerate(node):
                item_path = path + [f"[{idx}]"]
                yield from iter_answerable_nodes(item, item_path)

    # --- Case 3: primitive at root (unlikely but safe) ---
    else:
        if path:
            yield {
                "path": path,
                "heading": path[-1],
                "requirements": [str(node)]
            }

async def run_agent(rfp_file: str, output_file: str = "rfp_response.json"):
    # 1. Initialize RAG and Reset DB
    rag_manager = RAGManager()
    rag_manager.reset_database()

    # 2. Initialize Chat Client
    chat_client = OpenAIChatClient(
      # model_id="qwen3:8b-40k",
        model_id="qwen3:14b-40k",
        api_key="ollama",
        base_url="http://localhost:11434/v1",
    )

    # 3. Define Tools
    docling_tool = MCPStreamableHTTPTool(
        name="docling_tool",
        description="Convert a PDF to Markdown using docling MCP server",
        url="http://localhost:8000/mcp",
    )

    print(f"RFP Target: {rfp_file}")
    print(f"Output File: {output_file}\n")

    # 4. Create Agent (EXTRACTION AGENT)
    agent = chat_client.create_agent(
        name="RFP Response Architect",
        instructions=f"""
        You are an expert Bid Manager. Your job is to extract all requirements, scope of work, response information and questions from an RFP..

        TOOLS
        -----
        You have access to three important tools to convert documents (docling_tool) and manage the knowledge base (rag_manager.add_document and rag_manager.query_knowledge):

        1. docling_tool
           - Provides access to a Docling MCP server that allows you to convert documents, e.g. PDF files to markdown.
           - Always use this to read the source PDFs to add to the RAG knowledge base. Do not guess their contents.

        2. rag_manager.add_document
           - Saves markdown text into the knowledge base.
           - Parameters:
               - content: the markdown text you want to store.
               - source: "RFP".
           - Always call this after using docling_tool to ingest the document.

        3. rag_manager.query_knowledge
           - Retrieves relevant text chunks from the knowledge base using semantic search.
           - Parameters:
               - query: Wwhat you are looking for.
               - source_filter: Use "RFP" when searching for the exact RFP requirement text.
           - You must use this tool whenever you need factual details from the RFP. Do not rely on your memory.

        REQUIREMENT EXTRACTION
        ----------------------
        You must:
        - Work through the RFP and extract the Scope of work, proposal response format and requirements, evaluation criteria and anything else that looks like we need to respond to.
        - Show the requirements from the RFP as they appea rin the RFP, do not intepret them, show them exactly as they appear in the RFP.

        OUTPUT FORMAT
        -------------
        Output the requirements in a JSON format, retaining the RFP document hierarchy exactly as it appears in the RFP - section by section.

        GUARDRAILS
        ----------
        You must use the following guardrails when developing a respone:
        - Never fabricate anything that is not supported by retrieved information from the knowledge base.
        - If information is incomplete, say so clearly and suggest that the bid team add more detail.
        - Use professional, confident, and concise language suitable for a formal RFP.

        Your goal is to extract the requirements from the RFP that is fully traceable back to the source documents.
        """,
        tools=[docling_tool, rag_manager.add_document, rag_manager.query_knowledge]
    )

    # 5. Run the Agent with a Phased Prompt
    try:
        result = await agent.run(
        f"""
        Use the workflow and tool usage described in your system instructions.

        INPUT FILES
        -----------
        - RFP-PDF: '{rfp_file}'

        TASK
        ----
        1. Perform your full ingestion phase for the RFP PDF:
        2. List all the of information you have extracted from the RFP PDF.

        The final output of this will be in a format that will be saved in a JSON document.
        """
        )

        print(f"Agent: {result}\n\n")

        return result
    finally:
        # Ensure proper cleanup of agent resources
        if hasattr(agent, 'cleanup'):
            await agent.cleanup()
        # Give async generators time to clean up
        await asyncio.sleep(0.1)


# --- NEW: Per-section response generation ---

async def generate_section_responses(requirements_json_path: str,
                                     answers_output_path: str = "rfp_answers.json"):
    """
    Loads the extracted RFP requirements JSON and has an Agent respond to each
    answerable JSON structure (e.g., 1.2.1, 1.2.2) one at a time.

    Traversal is bottom-up within each dictionary: last key first.
    Final output is a JSON file mapping section paths to answers.
    """

    # Reconnect to the existing RAG DB (which should contain the RFP text)
    rag_manager = RAGManager()

    chat_client = OpenAIChatClient(
        model_id="qwen3:14b-40k",
        api_key="ollama",
        base_url="http://localhost:11434/v1",
    )

    # Agent focused on writing responses, not extraction
    response_agent = chat_client.create_agent(
        name="RFP Response Writer",
        instructions="""
        You are an expert RFP response writer for a Managed IT Services provider.

        TASK
        ----
        You will be given:
        - A section path (e.g., "RFP > 1. SCOPE OF WORK > 1.2 Scope > 1.2.1 Network monitoring and optimization")
        - A section id (e.g., "1.2.1")
        - A section title
        - The original list of RFP requirements (bullets)

        Your job:
        - Draft a concise, professional response in Markdown.
        - Address each requirement explicitly.
        - Align your answer with standard managed IT services capabilities.
        - Use confident but not over-promising language.
        - Do not include JSON or code fences; output plain Markdown only.

        USE OF KNOWLEDGE BASE
        ---------------------
        - Use rag_manager.query_knowledge when you need to refer back to the exact
          RFP wording or surrounding context.
        - Summarize, but do not contradict, the RFP requirements.

        STYLE
        -----
        - Formal, clear, and easy to score by an evaluator.
        - Use short paragraphs and bullet points where appropriate.
        """,
        tools=[rag_manager.query_knowledge]
    )

    print(f"Loading requirements from: {requirements_json_path}")
    with open(requirements_json_path, "r", encoding="utf-8") as f:
        req_data = json.load(f)

    # Handle either {"RFP": {...}} or a direct hierarchy root
    rfp_tree = req_data.get("RFP", req_data)

    answers = {}
    try:
        for item in iter_answerable_nodes(rfp_tree):
            path = item["path"]
            heading = item["heading"]
            requirements = item["requirements"]

            section_path = " > ".join(path)
            parts = heading.split()
            if parts and re.match(r"^\d+(\.\d+)*$", parts[0]):
                section_id = parts[0]
                section_title = " ".join(parts[1:]) or heading
            else:
                section_id = ""
                section_title = heading

            bullets_md = "\n".join(f"- {b}" for b in requirements)

            print(f"\n--- Generating answer for: {section_path} ---")

            prompt = f"""
You are preparing a response to an RFP section.

RFP section path:
{section_path}

Section ID: {section_id}
Section Title: {section_title}

Original RFP requirements:
{bullets_md}

TASK
----
1. If helpful, first call your knowledge tool to retrieve relevant RFP context.
2. Then draft a concise, professional answer in Markdown that:
   - Clearly addresses each bullet above.
   - Uses first-person plural ("we") to describe the service provider.
   - Is suitable to paste directly into an RFP response document.

Remember: output Markdown only, no JSON, no code fences.
"""

            result = await response_agent.run(prompt)
            answer_text = extract_main_content(result).strip()

            answers[section_path] = {
                "section_id": section_id,
                "section_title": section_title,
                "requirements": requirements,
                "answer_markdown": answer_text
            }

    finally:
        if hasattr(response_agent, 'cleanup'):
            await response_agent.cleanup()
        await asyncio.sleep(0.1)

    output_obj = {
        "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "source_requirements_file": requirements_json_path,
        "answers": answers
    }

    with open(answers_output_path, "w", encoding="utf-8") as f:
        json.dump(output_obj, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Section responses saved to: {answers_output_path}")


def _render_markdown_to_docx(doc: Document, md_text: str):
    """
    Very simple Markdown-to-Word renderer:
      - # / ## / ### headings -> Word headings
      - - / * bullets -> bullet list
      - 1. / 2. numbered lines -> numbered list
      - everything else -> normal paragraph

    This is intentionally minimal but good enough for typical RFP responses.
    """
    lines = md_text.splitlines()

    for line in lines:
        raw = line.rstrip("\n")
        text = raw.strip()

        if not text:
            # Blank line -> spacing paragraph
            doc.add_paragraph("")
            continue

        # Headings
        if text.startswith("### "):
            doc.add_heading(text[4:], level=4)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=3)
        elif text.startswith("# "):
            doc.add_heading(text[2:], level=2)

        # Bullets
        elif text.startswith("- ") or text.startswith("* "):
            doc.add_paragraph(text[2:], style="List Bullet")

        # Numbered list (e.g. "1. Item")
        elif re.match(r"^\d+\.\s+", text):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", text), style="List Number")

        else:
            # Plain paragraph
            doc.add_paragraph(text)


def answers_json_to_word(answers_json_path: str, docx_path: str):
    """
    Converts the section responses JSON (answers_file) into a Word document.

    Structure:
      - Top-level title: "RFP Section Responses"
      - For each section:
          Heading: "<section_id> <section_title>" (or the section_path if no ID)
          Subheading: "Original Requirements"
          Bullet list of requirements
          Subheading: "Response"
          Rendered Markdown answer
    """

    print(f"Converting section responses to Word: {answers_json_path} -> {docx_path}")

    with open(answers_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    answers = data.get("answers", {})

    doc = Document()
    doc.add_heading("RFP Section Responses", level=1)

    # Helper to derive a sort key using section_id if available
    def sort_key(item):
        section_path, section_data = item
        section_id = section_data.get("section_id", "")
        if section_id:
            # Try to sort numerically by section_id parts: "1.2.10" -> [1,2,10]
            parts = []
            for p in section_id.split("."):
                try:
                    parts.append(int(p))
                except ValueError:
                    # Fallback if something non-numeric sneaks in
                    parts.append(p)
            return (parts, section_path)
        else:
            return ([9999], section_path)  # push unknown IDs to the end

    # Iterate in a stable, roughly RFP-like order
    for section_path, section_data in sorted(answers.items(), key=sort_key):
        section_id = section_data.get("section_id", "")
        section_title = section_data.get("section_title", "")
        requirements = section_data.get("requirements", [])
        answer_md = section_data.get("answer_markdown", "")

        # Section heading
        if section_id:
            heading_text = f"{section_id} {section_title}".strip()
        else:
            # Fallback to full path if no explicit ID
            heading_text = section_path

        doc.add_heading(heading_text, level=2)

        # Original requirements
        if requirements:
            doc.add_heading("Original Requirements", level=3)
            for req in requirements:
                doc.add_paragraph(str(req), style="List Bullet")

        # Response
        if answer_md.strip():
            doc.add_heading("Response", level=3)
            _render_markdown_to_docx(doc, answer_md)

    doc.save(docx_path)
    print(f"✓ Section responses Word document generated: {docx_path}")


if __name__ == "__main__":
    # Default file paths
    default_rfp = "/home/ubuntu/MAF/Sample-RFP-Managed-Services.pdf"
    default_output = "/home/ubuntu/MAF/rfp_requirements.json"

    # Parse command line arguments
    if len(sys.argv) >= 2:
        rfp_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) >= 3 else default_output
    else:
        print("Usage: python rfp_agent.py <rfp_pdf> [output_json]")
        print(f"\nUsing default files:")
        print(f"  RFP: {default_rfp}")
        print(f"  Output: {default_output}\n")
        rfp_file = default_rfp
        output_file = default_output

    # 1) Run extraction agent and save result
    result = asyncio.run(run_agent(rfp_file, output_file))
    save_to_json(result, output_file)

    # 2) Run per-section response agent over the extracted JSON
    base, ext = os.path.splitext(output_file)
    answers_file = f"{base}_answers.json"
    asyncio.run(generate_section_responses(output_file, answers_file))

    # 3) Convert section responses JSON to Word
    answers_docx = f"{base}_answers.docx"
    answers_json_to_word(answers_file, answers_docx)
