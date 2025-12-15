## Author: Steve Harris
# Purpose: MAF Agent with RAG that checks Company Info against an RFP
# NOTES:
#  - Requires: pip install chromadb python-docx

import asyncio
import uuid
import sys
import chromadb
from chromadb.utils import embedding_functions
from agent_framework import ChatAgent, MCPStreamableHTTPTool
from agent_framework.openai import OpenAIChatClient
from docx import Document
from datetime import datetime
import logging
import os, re, json

# Suppress noisy MCP async generator cleanup warnings
logging.getLogger('asyncio').setLevel(logging.CRITICAL)

# --- RAG Manager Class Implementation ---

class RAGManager:
    def __init__(self):
        self.client = chromadb.Client()
        self.ef = embedding_functions.DefaultEmbeddingFunction()
        self.collection = self.client.get_or_create_collection(
            name="rfp_knowledge",
            embedding_function=self.ef
        )

    def reset_database(self):
        """Clears existing data to ensure fresh context."""
        try:
            self.client.delete_collection("rfp_knowledge")
        except:
            pass
        self.collection = self.client.get_or_create_collection(
            name="rfp_knowledge",
            embedding_function=self.ef
        )
        print(f"Database reset successfully.\n")
        return "Database reset successfully."

    def add_document(self, content: str, source: str) -> str:
        """
        Adds text content to the vector database for later retrieval when answering questions.
        Args:
            content: The text (markdown) to add.
            source: A label for where this text came from (e.g., "Company Profile", "RFP").
        """

        # Split by double newline to create logical chunks

        chunks = [c for c in content.split("\n\n") if c.strip()]

        if not chunks:
            return "No content found to add."

        ids = [str(uuid.uuid4()) for _ in chunks]
        metadatas = [{"source": source} for _ in chunks]

        self.collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        print(f"--- [RAG] Ingested {len(chunks)} chunks from source: {source} ---")
        return f"Successfully added {len(chunks)} chunks from '{source}' to the knowledge base."

    def query_knowledge(self, query: str, n_results: int = 10, source_filter: str = None) -> str:
        """
        Queries the vector database for relevant context.
        Args:
            query: What you are looking for
            n_results: Number of results to return (default: 10)
            source_filter: Optional source to filter results by (e.g., "Company-Info", "RFP")
        """

        # VERIFY: Check what's being queried
        print(f"--- [RAG] Querying: '{query}' with filter: {source_filter} ---")

        query_params = {
            "query_texts": [query],
            "n_results": n_results
        }

        # Add metadata filter if source_filter is provided
        if source_filter:
            query_params["where"] = {"source": source_filter}

        results = self.collection.query(**query_params)

        context = ""
        if results['documents']:
            for i, (doc, meta) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
                source_label = meta.get('source', 'Unknown')
                context += f"-- Result {i+1} (Source: {source_label}) --\n{doc}\n\n"

        # VERIFY: Check results
        if results['documents'] and results['documents'][0]:
             print(f"--- [RAG] Found {len(results['documents'][0])} results ---")
        else:
             print(f"--- [RAG] NO RESULTS FOUND ---")

        print(f"Knowledge query executed: '{query}' (filter: {source_filter})\n")
        return context if context else "No relevant information found in the knowledge base."


def save_to_json(result, output_filename: str = "rfp_response.json"):
    """Extracts JSON payload for output file and saves model reasoning trace separately for audit."""

    # Extract text from Agent response
    if hasattr(result, 'text'):
        content = result.text
    elif hasattr(result, 'content'):
        content = result.content
    else:
        content = str(result)

    # Ensure audit filename ends in .audit.json
    base, ext = os.path.splitext(output_filename)
    audit_filename = f"{base}.audit.json"

    # --- Extract <think> reasoning content for audit log ---
    think_matches = re.findall(r"<think>(.*?)</think>", content, flags=re.DOTALL)
    reasoning_log = [t.strip() for t in think_matches if t.strip()]

    # --- Remove <think> blocks from main output content ---
    content_no_think = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()

    # --- Extract JSON fenced block if present ---
    json_str = None
    fence = "```json"
    if fence in content_no_think:
        start = content_no_think.index(fence) + len(fence)
        end = content_no_think.find("```", start)
        if end == -1:
            end = len(content_no_think)
        json_str = content_no_think[start:end].strip()
    else:
        # fallback: find braces
        first_brace = content_no_think.find("{")
        last_brace = content_no_think.rfind("}")
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_str = content_no_think[first_brace:last_brace + 1].strip()

    # --- Parse JSON or fallback ---
    if json_str:
        try:
            data = json.loads(json_str)
        except Exception:
            data = {
                "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "raw_output": content_no_think
            }
    else:
        data = {
            "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "raw_output": content_no_think
        }

    # --- Write main output JSON ---
    with open(output_filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # --- Write reasoning audit log ---
    audit_record = {
        "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "reasoning_trace": reasoning_log,
        "raw_full_response": content  # full record for evidence
    }
    with open(audit_filename, "w", encoding="utf-8") as f:
        json.dump(audit_record, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Extracted JSON saved to: {output_filename}")
    print(f"✓ Full audit + reasoning saved to: {audit_filename}")

    return output_filename, audit_filename


# --- Helper: strip <think> blocks for per-section answers ---

def extract_main_content(result) -> str:
    """
    Returns the model's main content with <think> blocks stripped out.
    This is used for the per-section answer agent.
    """
    if hasattr(result, 'text'):
        content = result.text
    elif hasattr(result, 'content'):
        content = result.content
    else:
        content = str(result)

    content_no_think = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
    return content_no_think


# --- Helper: iterate answerable nodes (bottom-up within each dict) ---

def iter_answerable_nodes(node, path=None):
    """
    Recursively walks the RFP requirements JSON and yields
    answerable nodes for ALL layers of the JSON.

    Rules:
      - If we see a dict: iterate keys in reverse order (bottom-up within section).
      - If we see a list of primitives: treat the list as a single answerable node.
      - If we see a list of dicts/lists: recurse into each item.
      - If we see a single primitive (string/number/bool/None): treat it as
        an answerable node with a single requirement.

    Yields dicts with:
      - path: list of keys/indexes from root to this node
      - heading: the "label" for this node (usually the last key or index)
      - requirements: list of strings (the bullets or text to respond to)
    """
    if path is None:
        path = []

    # Helper to decide if something is a "primitive"
    def is_primitive(x):
        return isinstance(x, (str, int, float, bool)) or x is None

    # --- Case 1: dict ---
    if isinstance(node, dict):
        # Bottom-up within each section
        for key in reversed(list(node.keys())):
            value = node[key]
            new_path = path + [key]

            # If the value is a list, decide whether it's a leaf or container
            if isinstance(value, list):
                if all(is_primitive(v) for v in value):
                    # Leaf: list of primitive requirements
                    yield {
                        "path": new_path,
                        "heading": key,
                        "requirements": [str(v) for v in value]
                    }
                else:
                    # Container: recurse into each element
                    for idx, item in enumerate(value):
                        item_path = new_path + [f"[{idx}]"]
                        yield from iter_answerable_nodes(item, item_path)

            # If the value is another dict, recurse into it
            elif isinstance(value, dict):
                yield from iter_answerable_nodes(value, new_path)

            # If the value is a single primitive, treat it as a leaf requirement
            elif is_primitive(value):
                yield {
                    "path": new_path,
                    "heading": key,
                    "requirements": [str(value)]
                }

    # --- Case 2: list at this level ---
    elif isinstance(node, list):
        if node and all(is_primitive(v) for v in node):
            # Leaf list of primitives
            heading = path[-1] if path else "[root_list]"
            yield {
                "path": path,
                "heading": heading,
                "requirements": [str(v) for v in node]
            }
        else:
            # Container list: recurse into children
            for idx, item in enumerate(node):
                item_path = path + [f"[{idx}]"]
                yield from iter_answerable_nodes(item, item_path)

    # --- Case 3: primitive at root (unlikely but safe) ---
    else:
        if path:
            yield {
                "path": path,
                "heading": path[-1],
                "requirements": [str(node)]
            }

async def run_agent(rfp_file: str, company_info_file: str, output_file: str = "rfp_response.json"):
    # 1. Initialize RAG DB
    rag_manager.reset_database()

    # 2. Initialize Chat Client
    chat_client = OpenAIChatClient(
        model_id="qwen3:14b-40k",
        api_key="ollama",
        base_url="http://localhost:11434/v1",
    )

    # 3. Define Tools
    docling_tool = MCPStreamableHTTPTool(
        name="docling_tool",
        description="Use for document conversion. Convert a PDF to Markdown using docling MCP server.",
        url="http://localhost:8000/mcp",
    )

    print(f"RFP Target: {rfp_file}")
    print(f"Company Info: {company_info_file}")
    print(f"Output File: {output_file}\n")

    # ============================================================
    # NEW: EXPLICIT COMPANY INFO INGESTION BEFORE AGENT RUNS
    # ============================================================
    if company_info_file and os.path.exists(company_info_file):
        print(f"\n{'='*60}")
        print(f"PRE-INGESTING COMPANY INFORMATION")
        print(f"{'='*60}\n")

        # Create a temporary agent just for ingestion
        ingestion_agent = chat_client.create_agent(
            name="Document Ingestion Agent",
            instructions="""
            You are a document ingestion specialist.

            Your ONLY job is to:
            1. Use docling_tool to convert the provided PDF to markdown
            2. Call rag_manager.add_document with the markdown content and source="Company-Info"

            Do NOT do anything else. Do NOT analyze or summarize.
            Just convert and ingest.
            """,
            tools=[docling_tool, rag_manager.add_document]
        )

        try:
            ingestion_result = await ingestion_agent.run(
                f"Convert '{company_info_file}' to markdown using docling_tool, "
                f"then save it to the knowledge base with source='Company-Info'."
            )
            print(f"\nIngestion Result: {ingestion_result}\n")

            # VERIFY ingestion worked
            collection_count = rag_manager.collection.count()
            print(f"✓ Knowledge base now contains {collection_count} chunks\n")

            if collection_count == 0:
                print("⚠️  WARNING: No documents were ingested! Proceeding anyway...")

        except Exception as e:
            print(f"❌ Error during company info ingestion: {e}")
            print("Proceeding without company information...\n")
        finally:
            if hasattr(ingestion_agent, 'cleanup'):
                await ingestion_agent.cleanup()
            await asyncio.sleep(0.1)
    else:
        print(f"⚠️  Company info file not found or not provided: {company_info_file}")
        print("Proceeding without company information...\n")

    # ============================================================
    # NOW CREATE THE EXTRACTION AGENT (focused only on RFP)
    # ============================================================
    agent = chat_client.create_agent(
        name="RFP Requirements Extractor",
        instructions=f"""
        You are an expert Bid Manager focused on extracting requirements from RFPs.

        TOOLS
        -----
        You have access to:

        1. docling_tool
           - Converts PDF documents to markdown
           - Use this to read the RFP file
           - Do NOT pass numeric values for max_size

        2. rag_manager.query_knowledge
           - The knowledge base has already been populated with company information
           - You can query it if you need to reference company capabilities
           - Use source_filter="Company-Info" when querying

        REQUIREMENT EXTRACTION
        ----------------------
        Your job is to:
        - Convert the RFP PDF to markdown using docling_tool
        - Extract the Scope of Work, proposal response format, requirements, evaluation criteria and any other RFP sections that need a response.
        - Show requirements exactly as they appear in the RFP
        - Do NOT interpret or modify the requirements

        OUTPUT FORMAT
        -------------
        Output the requirements in JSON format, retaining the RFP document hierarchy exactly.

        NOTE: Company information has ALREADY been ingested into the knowledge base.
        You do NOT need to ingest it again. Focus only on extracting RFP requirements.
        """,
        tools=[docling_tool, rag_manager.query_knowledge]  # Note: no add_document needed
    )

    # 5. Run the Agent (now focused only on RFP extraction)
    try:
        result = await agent.run(
            f"""
            Extract all requirements from the RFP file: '{rfp_file}'

            1. Use docling_tool to convert the RFP to markdown
            2. Extract and structure all requirements in JSON format
            3. Retain the original RFP hierarchy

            The knowledge base already contains company information if you need to reference it.
            """
        )

        print(f"Agent: {result}\n\n")
        return result

    finally:
        if hasattr(agent, 'cleanup'):
            await agent.cleanup()
        await asyncio.sleep(0.1)

# --- Per-section response generation with optional company info RAG ---

async def generate_section_responses(requirements_json_path: str,
                                     answers_output_path: str = "rfp_answers.json",
                                     company_info_path: str = "MyCompany-Capabilities-Test.pdf"):  # NEW arg
    """
    Loads the extracted RFP requirements JSON and has an Agent respond to each answerable JSON structure (e.g., 1.2.1, 1.2.2) one at a time.

    Traversal is bottom-up within each dictionary: last key first.
    Final output is a JSON file mapping section paths to answers.

    Assumes the RAG database has already been populated with Company information (source='Company-Info', if provided) by the extraction agent.
    """

    chat_client = OpenAIChatClient(
        model_id="qwen3:14b-40k",
        api_key="ollama",
        base_url="http://localhost:11434/v1",
    )

    # Agent focused on writing responses, not extraction
    response_agent = chat_client.create_agent(
        name="RFP Response Writer",
        instructions="""
        You are an expert RFP response writer for a Managed IT Services provider.

        TASK
        ----
        You will be given:
        - A section path (e.g., "RFP > 1. SCOPE OF WORK > 1.2 Scope > 1.2.1 Network monitoring and optimization")
        - A section id (e.g., "1.2.1")
        - A section title
        - The original list of RFP requirements (bullets)

        Your job:
        - Draft a concise, professional response in Markdown.
        - Address each requirement explicitly.
        - Align your answer with standard managed IT services capabilities provided by the Company..
        - Extract When company information is available in the knowledge base, explain how our specific capabilities, services, experience, certifications, and differentiators meet these requirements.
        - Do not include JSON or code fences; output plain Markdown only.

        USE OF KNOWLEDGE BASE
        ---------------------
        - rag_manager.query_knowledge retrieves the company information from the knwoledge base.
        - Call it with query and source_filter="Company-Info" when you need company capabilities.
        - Call this tool whenever you need factual details about the company's capabilities, services, experience, certifications, locations, and differentiators.
        - Prefer retrieved company-specific details over generic statements.
        - Never fabricate capabilities or claims that are not supported by retrieved information.
        - Summarize, but do not contradict, the RFP requirements or company information.

        STYLE
        -----
        - Formal, clear, and easy to score by an evaluator.
        - Use first-person plural ("we") for the service provider.
        - Use short paragraphs and bullet points where appropriate.
        """,
        tools=[rag_manager.query_knowledge]
    )

    print(f"Loading requirements from: {requirements_json_path}")
    with open(requirements_json_path, "r", encoding="utf-8") as f:
        req_data = json.load(f)

    # Handle either {"RFP": {...}} or a direct hierarchy root
    rfp_tree = req_data.get("RFP", req_data)

    answers = {}
    try:
        for item in iter_answerable_nodes(rfp_tree):
            path = item["path"]
            heading = item["heading"]
            requirements = item["requirements"]

            section_path = " > ".join(path)
            parts = heading.split()
            if parts and re.match(r"^\d+(\.\d+)*$", parts[0]):
                section_id = parts[0]
                section_title = " ".join(parts[1:]) or heading
            else:
                section_id = ""
                section_title = heading

            bullets_md = "\n".join(f"- {b}" for b in requirements)

            print(f"\n--- Generating answer for: {section_path} ---")

            # UPDATED prompt to explicitly mention company info + knowledge queries
            prompt = f"""
            You are preparing a response to an RFP section.

            RFP section path:
            {section_path}

            Section ID: {section_id}
            Section Title: {section_title}

            Original RFP requirements:
            {bullets_md}

            TASK
            ----
            1. Call your knowledge tool (rag_manager.query_knowledge) using a query that includes the section title and key requirement terms. Use this to retrieve any relevant company information (from 'Company-Info') that describes our capabilities, services, experience, certifications, and differentiators related to this section.

            2. Then draft a concise, professional answer in Markdown that:
               - Clearly addresses each bullet above.
               - Uses first-person plural ("we") to describe the service provider.
               - Incorporates retrieved company-specific details wherever relevant, instead of generic statements.
               - Is suitable to paste directly into an RFP response document.

            Remember: output Markdown only, no JSON, no code fences.
            """

            result = await response_agent.run(prompt)
            answer_text = extract_main_content(result).strip()
            print(f"Answer Text: {answer_text}\n")

            answers[section_path] = {
                "section_id": section_id,
                "section_title": section_title,
                "requirements": requirements,
                "answer_markdown": answer_text
            }

    finally:
        if hasattr(response_agent, 'cleanup'):
            await response_agent.cleanup()
        await asyncio.sleep(0.1)

    output_obj = {
        "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "source_requirements_file": requirements_json_path,
        "answers": answers
    }

    with open(answers_output_path, "w", encoding="utf-8") as f:
        json.dump(output_obj, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Section responses saved to: {answers_output_path}")


def _render_markdown_to_docx(doc: Document, md_text: str):
    """
    Very simple Markdown-to-Word renderer:
      - # / ## / ### headings -> Word headings
      - - / * bullets -> bullet list
      - 1. / 2. numbered lines -> numbered list
      - everything else -> normal paragraph

    This is intentionally minimal but good enough for typical RFP responses.
    """
    lines = md_text.splitlines()

    for line in lines:
        raw = line.rstrip("\n")
        text = raw.strip()

        if not text:
            # Blank line -> spacing paragraph
            doc.add_paragraph("")
            continue

        # Headings
        if text.startswith("### "):
            doc.add_heading(text[4:], level=4)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=3)
        elif text.startswith("# "):
            doc.add_heading(text[2:], level=2)

        # Bullets
        elif text.startswith("- ") or text.startswith("* "):
            doc.add_paragraph(text[2:], style="List Bullet")

        # Numbered list (e.g. "1. Item")
        elif re.match(r"^\d+\.\s+", text):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", text), style="List Number")

        else:
            # Plain paragraph
            doc.add_paragraph(text)


def answers_json_to_word(answers_json_path: str, docx_path: str):
    """
    Converts the section responses JSON (answers_file) into a Word document.

    Structure:
      - Top-level title: "RFP Section Responses"
      - For each section:
          Heading: "<section_id> <section_title>" (or the section_path if no ID)
          Subheading: "Original Requirements"
          Bullet list of requirements
          Subheading: "Response"
          Rendered Markdown answer
    """

    print(f"Converting section responses to Word: {answers_json_path} -> {docx_path}")

    with open(answers_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    answers = data.get("answers", {})

    doc = Document()
    doc.add_heading("RFP Section Responses", level=1)

    # Helper to derive a sort key using section_id if available
    def sort_key(item):
        section_path, section_data = item
        section_id = section_data.get("section_id", "")
        if section_id:
            # Try to sort numerically by section_id parts: "1.2.10" -> [1,2,10]
            parts = []
            for p in section_id.split("."):
                try:
                    parts.append(int(p))
                except ValueError:
                    # Fallback if something non-numeric sneaks in
                    parts.append(p)
            return (parts, section_path)
        else:
            return ([9999], section_path)  # push unknown IDs to the end

    # Iterate in a stable, roughly RFP-like order
    for section_path, section_data in sorted(answers.items(), key=sort_key):
        section_id = section_data.get("section_id", "")
        section_title = section_data.get("section_title", "")
        requirements = section_data.get("requirements", [])
        answer_md = section_data.get("answer_markdown", "")

        # Section heading
        if section_id:
            heading_text = f"{section_id} {section_title}".strip()
        else:
            # Fallback to full path if no explicit ID
            heading_text = section_path

        doc.add_heading(heading_text, level=2)

        # Original requirements
        if requirements:
            doc.add_heading("Original Requirements", level=3)
            for req in requirements:
                doc.add_paragraph(str(req), style="List Bullet")

        # Response
        if answer_md.strip():
            doc.add_heading("Response", level=3)
            _render_markdown_to_docx(doc, answer_md)

    doc.save(docx_path)
    print(f"✓ Section responses Word document generated: {docx_path}")


if __name__ == "__main__":
    # Default file paths
    default_rfp = "/home/ubuntu/MAF/Sample-RFP-Managed-Services.pdf"
    default_output = "/home/ubuntu/MAF/rfp_requirements.json"
    default_company_info = "/home/ubuntu/MAF/MyCompany-Capabilities-Test.pdf"

    # Instantiate a global rag_manager
    rag_manager = RAGManager()

    # Parse command line arguments
    if len(sys.argv) >= 2:
        rfp_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) >= 3 else default_output
        company_info_file = sys.argv[3] if len(sys.argv) >= 4 else default_company_info  # NEW
    else:
        print("Usage: python python_script.py <rfp_pdf> [output_json] [company_info_file]")
        print(f"\nUsing default files:")
        print(f"  RFP: {default_rfp}")
        print(f"  Output: {default_output}")
        print(f"  Company Info: {default_company_info}\n")
        rfp_file = default_rfp
        output_file = default_output
        company_info_file = default_company_info

    # 1) Run extraction agent and save result
    result = asyncio.run(run_agent(rfp_file, company_info_file, output_file))
    save_to_json(result, output_file)

    # 2) Run per-section response agent over the extracted JSON (with optional company info)
    base, ext = os.path.splitext(output_file)
    answers_file = f"{base}_answers.json"
    asyncio.run(generate_section_responses(output_file, answers_file, company_info_file))  # UPDATED

    # 3) Convert section responses JSON to Word
    answers_docx = f"{base}_answers.docx"
    answers_json_to_word(answers_file, answers_docx)

